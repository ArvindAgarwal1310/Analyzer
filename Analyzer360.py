import docx
import pandas
import docx2pdf
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import smtplib
from email.message import EmailMessage
from docx.enum.section import WD_SECTION_START
from docx import *
import matplotlib.pyplot as mpt
df=pandas.read_csv("football_players.csv")
choice_pdf=1
#df.dropna(inplace=True)
#for a in (df.columns):
#    print(a)
docuu=docx.Document()
#docuu.sections[0].page_height=docx.shared.Inches(500)
shd = OxmlElement("w:background")
shd.set(qn("w:color"), "#B2DFEE")
docuu.element.insert(0,shd)
shd1 = OxmlElement("w:displayBackgroundShape")
docuu.settings.element.insert(0,shd1)
#new_section = docuu.add_section(start_type=WD_SECTION_START.CONTINUOUS)
#new_section.start_type = WD_SECTION_START.CONTINUOUS
docuu.add_picture("Analyser-360logo.jpg",width=docx.shared.Inches(2))
ns=docuu.styles["Quote"]
docuu.add_paragraph(style=ns)
docuu.add_heading("Analyser-360",0)
df=df.replace("yes",True)
df=df.replace("Yes",True)
df=df.replace("no",False)
df=df.replace("No",False)
Senders_list=["arvindagarwal839@gmail.com","aagarwal1310@gmail.com"]
'''if (str(df.iloc[-1][1]).lower()=="pdf"):
    choice_pdf=1
df.drop(df.tail(3).index,inplace=True)'''
def func(pct, allvalues):
    absolute = int(pct / 100.*sum(allvalues))
    return "{:.2f}%\n({:d})".format(pct, absolute)
for a in df.select_dtypes(include=object):
    if("NetWorth" in a or "worth" in a or "Worth" in a):
        df[a].replace("$","")
        df[a]=df[a].replace(",","")
for a in df.select_dtypes(include=bool):
    df[a].fillna(False)
    description=df[a].describe()
    #print(description)
    ll=[description["freq"],description["count"]-description["freq"]]
    #mexplodes=[0.1,0,0]
    mpt.pie(ll,labels=[description["top"],"False" if description["top"]==True else "True"],autopct = lambda pct: func(pct, ll))
    mpt.legend(title=str(a).capitalize().replace("_"," "))
    #mpt.title(str(a).capitalize())
    mpt.savefig("imaggggggeee.png", transparent=True)
    docuu.add_heading(str(a).replace("_", " ").capitalize(), 1)
    docuu.add_picture("imaggggggeee.png", width=docx.shared.Inches(6.5), height=docx.shared.Inches(5))
    mpt.close()
    docuu.add_heading(" ", 0)
for a in df.select_dtypes(include=object,exclude=bool):
    if("id" in str(a).lower() or "roll" in str(a).lower() or "date" in str(a).lower() or "rank" in str(a).lower()):
        continue
    #print(f"dtype= {a}")
    #print(df[a].describe())
    #print(set(df[a]))
    ind_l=[]
    ind_v=[]
    str_form=str(df[a].to_string())
    for ind_val in set(df[a]):
        ind_l.append(str(ind_val))
        ind_v.append(int(str_form.count(str(ind_val))))
    docuu.add_heading(str(a).replace("_", " ").capitalize(), 1)

    if(len(ind_v)<8):
        mpt.bar(ind_l, ind_v)
        mpt.xlabel(str(a).capitalize().replace("_", " "))
        mpt.ylabel("Frequency")
        mpt.title(str(a).capitalize().replace("_", " "))
        # mpt.title(str(a).capitalize())
        sss ="imewwq.png"
        # docuu.add_picture(mpt)
        mpt.savefig(sss, transparent=True)
        docuu.add_picture(sss, width=docx.shared.Inches(4))
        #docuu.add_paragraph(str(df[a].describe()["mean":"max"]).capitalize()).italic = True
        # docuu.add_paragraph(str(df[a].std())).italic = True
        # docuu.add_paragraph(str(df[a].max())).italic = True
        mpt.close()
    docuu.add_paragraph(style="Heading 2")
    docuu.add_paragraph(f"Maximum occurence is of {df[a].describe().top}  Frequency : {df[a].describe().freq}")
    docuu.add_heading(" ", 0)
    #docuu.add_page_break()
for a in df.select_dtypes(include=int or float):
    if("id" in str(a).lower() or "roll" in str(a).lower() or "date" in str(a).lower() or "rank" in str(a).lower()):
        continue
    if(len(df[a].unique())<3):
        #print(a)
        continue
    df[a] = df[a].fillna(int(0))
    #print(df[a].describe(),"\n\n")
    ll=[df[a].mean(),df[a].std(),df[a].max()]
    pos_ll=[abs(df[a].mean()),abs(df[a].std()),abs(df[a].max())]
    mexplodes=[0.05,0,0]
    mpt.pie(pos_ll,labels=["mean","std","max"],explode=mexplodes,autopct = lambda pct: func(pct, ll))
    mpt.legend(title=str(a).capitalize().replace("_"," "))
    sss="inttim.png"
    mpt.savefig(sss,transparent=True)
    docuu.add_heading(str(a).replace("_"," ").capitalize(),1)
    docuu.add_picture(sss,width=docx.shared.Inches(6.5),height=docx.shared.Inches(5))
    docuu.add_paragraph(style="Heading 2")
    docuu.add_paragraph(str(df[a].describe()["mean":"max"]).capitalize()).italic=True
    mpt.close()
    docuu.add_heading(" ", 0)
new_section = docuu.add_section(start_type=WD_SECTION_START.NEW_PAGE)
new_section.start_type = WD_SECTION_START.NEW_PAGE
docuu.add_heading("Pro Analysis",1)
docuu.add_picture("Analyser-360logo.jpg",width=docx.shared.Inches(1))
file_wformat="analysed-360.docx"
docuu.save("analysed-360.docx")
print("Data Analysed.")
if(choice_pdf==1):
    docx2pdf.convert("analysed-360.docx","analysed-360.pdf")
    file_wformat="analysed-360.pdf"
e_addr='report.analyser360@gmail.com'
e_pass='juesceotigevavuz'
msg=EmailMessage()
msg['Subject']='Analysis by Analyser-360'
msg['From']=e_addr
msg['To']=Senders_list
msg.set_content(f'This email is auto generated.\nDocument attached below\n\n')
with open(file_wformat,'rb') as f:
    file_data=f.read()
    file_name=f.name
msg.add_attachment(file_data,maintype='application',subtype='octet-stream',filename=file_name)
with smtplib.SMTP('smtp.gmail.com',587) as smtp:
    smtp.ehlo()
    smtp.starttls()
    smtp.ehlo()
    smtp.login(e_addr,e_pass)
    smtp.send_message(msg)